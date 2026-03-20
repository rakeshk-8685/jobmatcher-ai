"""
Resume Routes
API endpoints for resume parsing
"""

from fastapi import APIRouter, UploadFile, File
from pydantic import BaseModel
from typing import Optional, List
import re
import io
from app.services.ats_scoring import _executor

router = APIRouter()

COMMON_SKILLS = [
    'python', 'javascript', 'typescript', 'react', 'angular', 'vue',
    'node.js', 'express', 'django', 'flask', 'java', 'c++', 'c#',
    'sql', 'mongodb', 'postgresql', 'mysql', 'redis', 'docker',
    'kubernetes', 'aws', 'azure', 'gcp', 'git', 'linux', 'rest api',
    'graphql', 'html', 'css', 'sass', 'tailwind', 'redux', 'next.js',
    'machine learning', 'data analysis', 'agile', 'scrum'
]

# Fast single-pass regex for skill extraction
SKILL_PATTERN = re.compile(
    r'\b(' + '|'.join(re.escape(s) for s in COMMON_SKILLS) + r')\b', 
    re.IGNORECASE
)

def extract_skills(text: str) -> list:
    """Extract skills from resume text using high-speed regex"""
    return list(set(SKILL_PATTERN.findall(text.lower())))

def extract_experience_years(text: str) -> int:
    """Extract years of experience from resume"""
    patterns = [
        r'(\d+)\+?\s*years?\s*(?:of\s*)?experience',
        r'(\d+)\+?\s*years?\s*(?:in|of)',
        r'experience[:\s]*(\d+)\+?\s*years?'
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text.lower())
        if match:
            return int(match.group(1))
    
    return 0

def extract_education(text: str) -> dict:
    """Extract education information"""
    education = {
        "degrees": [],
        "highest_level": None
    }
    
    degree_patterns = {
        'phd': ['ph.d', 'phd', 'doctorate'],
        'master': ['master', 'mba', 'm.s.', 'ms', 'm.a.'],
        'bachelor': ['bachelor', 'b.s.', 'bs', 'b.a.', 'ba'],
        'associate': ['associate', 'a.a.', 'a.s.']
    }
    
    text_lower = text.lower()
    
    for level, patterns in degree_patterns.items():
        for pattern in patterns:
            if pattern in text_lower:
                education["degrees"].append(level)
                if not education["highest_level"]:
                    education["highest_level"] = level
                break
    
    return education

def _parse_single_file_task(args):
    """Helper to parse a single file in a process pool (pickleable)"""
    filename, content_bytes = args
    filename = filename.lower()
    content = ""
    
    try:
        if filename.endswith('.pdf'):
            import fitz  # PyMuPDF (High performance C-based parsing)
            with fitz.open(stream=content_bytes, filetype="pdf") as doc:
                for page in doc:
                    content += page.get_text() + "\n"
                
        elif filename.endswith('.docx'):
            import docx
            docx_file = io.BytesIO(content_bytes)
            doc = docx.Document(docx_file)
            for para in doc.paragraphs:
                content += para.text + "\n"
                
        elif filename.endswith('.txt'):
            content = content_bytes.decode('utf-8', errors='ignore')
            
        else:
            return {"success": False, "error": "Unsupported file format", "filename": filename}
            
        # Clean up text
        content = re.sub(r'\s+', ' ', content).strip()
        
        # Extract metadata
        skills = extract_skills(content)
        experience_years = extract_experience_years(content)
        education = extract_education(content)
        
        # Extract email
        email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', content)
        email = email_match.group(0) if email_match else None
        
        # Extract phone
        phone_match = re.search(r'[\+]?[(]?[0-9]{3}[)]?[-\s\.]?[0-9]{3}[-\s\.]?[0-9]{4,6}', content)
        phone = phone_match.group(0) if phone_match else None
        
        return {
            "success": True,
            "filename": filename,
            "data": {
                "text": content,
                "skills": skills,
                "experience_years": experience_years,
                "education": education,
                "contact": {
                    "email": email,
                    "phone": phone
                }
            }
        }
    except Exception as e:
        return {"success": False, "error": str(e), "filename": filename}

@router.post("/batch-parse-resume-files")
def batch_parse_resume_files(files: List[UploadFile] = File(...)):
    """Parse multiple resume files in parallel using ProcessPool"""
    # Read all files into memory first (I/O bound)
    print(f"[AI Service] Incoming batch parse request for {len(files)} files")
    file_tasks = []
    for file in files:
        print(f"  - Reading file: {file.filename}")
        content_bytes = file.file.read()
        file_tasks.append((file.filename, content_bytes))
    
    # Process files in parallel (CPU bound)
    results = list(_executor.map(_parse_single_file_task, file_tasks))
    
    return {
        "success": True,
        "results": results
    }

@router.post("/parse-resume-file")
def parse_resume_file(file: UploadFile = File(...)):
    """Parse resume file (PDF/DOCX/TXT) and extract text"""
    print(f"[AI Service] Incoming single parse request for file: {file.filename}")
    content = ""
    filename = file.filename.lower()
    
    try:
        if filename.endswith('.pdf'):
            import fitz
            import io
            
            # Read PDF content
            pdf_content = file.file.read()
            with fitz.open(stream=pdf_content, filetype="pdf") as doc:
                for page in doc:
                    content += page.get_text() + "\n"
                
        elif filename.endswith('.docx'):
            import docx
            import io
            
            # Read DOCX content
            docx_content = file.file.read()
            docx_file = io.BytesIO(docx_content)
            doc = docx.Document(docx_file)
            
            for para in doc.paragraphs:
                content += para.text + "\n"
                
        elif filename.endswith('.txt'):
            content_bytes = file.file.read()
            content = content_bytes.decode('utf-8')
            
        else:
            return {"success": False, "error": "Unsupported file format"}
            
        # Clean up text
        content = re.sub(r'\s+', ' ', content).strip()
        
        # Extract metadata
        skills = extract_skills(content)
        experience_years = extract_experience_years(content)
        education = extract_education(content)
        
        # Extract email
        email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', content)
        email = email_match.group(0) if email_match else None
        
        # Extract phone
        phone_match = re.search(r'[\+]?[(]?[0-9]{3}[)]?[-\s\.]?[0-9]{3}[-\s\.]?[0-9]{4,6}', content)
        phone = phone_match.group(0) if phone_match else None
        
        return {
            "success": True,
            "data": {
                "text": content,
                "skills": skills,
                "experience_years": experience_years,
                "education": education,
                "contact": {
                    "email": email,
                    "phone": phone
                }
            }
        }
        
    except Exception as e:
        print(f"Error parsing file: {str(e)}")
        return {"success": False, "error": f"Failed to parse file: {str(e)}"}
